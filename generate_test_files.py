import os
import sys
import subprocess

# Ensure libraries are installed
def install_and_import(package_name, import_name=None):
    if import_name is None:
        import_name = package_name
    try:
        __import__(import_name)
    except ImportError:
        print(f"Installing {package_name}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])

install_and_import("python-docx", "docx")
install_and_import("reportlab")

import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def create_pdf_resume(path):
    print(f"Generating PDF Resume at: {path}")
    c = canvas.Canvas(path, pagesize=letter)
    c.setFont("Helvetica-Bold", 18)
    c.drawString(100, 750, "John Smith")
    c.setFont("Helvetica", 10)
    c.drawString(100, 735, "johnsmith@gmail.com | +91 9876543210")
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(100, 700, "Skills:")
    c.setFont("Helvetica", 10)
    c.drawString(100, 685, "Python, SQL, Power BI, Excel")
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(100, 650, "Education:")
    c.setFont("Helvetica", 10)
    c.drawString(100, 635, "B.Tech Computer Science, NIT Patna, 2024, CGPA: 8.5")
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(100, 600, "Experience:")
    c.setFont("Helvetica", 10)
    c.drawString(100, 585, "Data Analyst Intern - ABC Company, 6 months")
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(100, 550, "Projects:")
    c.setFont("Helvetica", 10)
    c.drawString(100, 535, "- Sales Dashboard using Power BI")
    c.drawString(100, 520, "- Customer Segmentation using Python")
    c.save()

def create_docx_resume(path):
    print(f"Generating DOCX Resume at: {path}")
    doc = docx.Document()
    
    doc.add_heading("Priya Sharma", 0)
    doc.add_paragraph("priya.sharma@email.com | +91 8765432109")
    
    doc.add_heading("Skills:", level=1)
    doc.add_paragraph("SQL, Excel, Tableau, Statistics")
    
    doc.add_heading("Education:", level=1)
    doc.add_paragraph("MBA Analytics, IIM Bangalore, 2023, CGPA: 9.0")
    
    doc.add_heading("Experience:", level=1)
    doc.add_paragraph("Business Analyst Intern - XYZ Corp, 1 year")
    
    doc.add_heading("Projects:", level=1)
    doc.add_paragraph("- Market Analysis Dashboard in Tableau")
    
    doc.save(path)

if __name__ == "__main__":
    uploads_dir = os.path.join(os.path.dirname(__file__), "backend", "uploads")
    os.makedirs(uploads_dir, exist_ok=True)
    
    pdf_path = os.path.join(uploads_dir, "john_smith.pdf")
    docx_path = os.path.join(uploads_dir, "priya_sharma.docx")
    
    create_pdf_resume(pdf_path)
    create_docx_resume(docx_path)
    
    # Also create copy in a root 'uploads' folder as requested
    root_uploads = os.path.join(os.path.dirname(__file__), "uploads")
    os.makedirs(root_uploads, exist_ok=True)
    
    import shutil
    shutil.copy2(pdf_path, os.path.join(root_uploads, "john_smith.pdf"))
    shutil.copy2(docx_path, os.path.join(root_uploads, "priya_sharma.docx"))
    
    print("\nSample resumes generated successfully!")
